from flask import Flask, render_template, request, jsonify, session, redirect, url_for, flash, send_file
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
from functools import wraps
from datetime import datetime, timedelta
import os
import json
import threading
import io
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.units import inch
from dotenv import load_dotenv
import time

# Import table generation utilities
from table_generator import TableDetector, HTMLTableGenerator
import markdown

# Markdown to HTML converter
def markdown_to_html(text):
    """Convert markdown text to HTML, including tables"""
    if not text:
        return ""
    
    html = markdown.markdown(
        text,
        extensions=['tables', 'fenced_code', 'nl2br']
    )
    return html

# Register as Jinja filter

# Load environment variables
load_dotenv()

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///research_agent.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'your-secret-key-change-in-production')
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=7)

db = SQLAlchemy(app)
app.jinja_env.filters['markdown'] = markdown_to_html
# ============================================================
# DATABASE MODELS
# ============================================================

class User(db.Model):
    """User model for authentication"""
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(120), unique=True, nullable=False)
    password = db.Column(db.String(255), nullable=False)
    is_admin = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    def set_password(self, password):
        self.password = generate_password_hash(password)
    
    def check_password(self, password):
        return check_password_hash(self.password, password)
    
    def __repr__(self):
        return f'<User {self.username}>'


class Query(db.Model):
    """Query model for storing research queries and responses"""
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    query_text = db.Column(db.Text, nullable=False)
    response = db.Column(db.Text)
    reasoning = db.Column(db.Text)
    status = db.Column(db.String(20), default='processing')
    task_type = db.Column(db.String(50))
    execution_time = db.Column(db.Float)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    tools_used = db.Column(db.String(255))
    error_message = db.Column(db.Text)
    
    # NEW FIELDS for table detection
    is_comparison_query = db.Column(db.Boolean, default=False)
    comparison_confidence = db.Column(db.Float, default=0.0)
    table_html = db.Column(db.Text)  # Store generated HTML table
    
    user = db.relationship('User', backref=db.backref('queries', lazy='dynamic'))
    
    def __repr__(self):
        return f'<Query {self.id}>'


class CompanyInfo(db.Model):
    """Company information model - stores all company knowledge"""
    __tablename__ = 'company_info'
    
    id = db.Column(db.Integer, primary_key=True)
    
    # Basic Info
    company_name = db.Column(db.String(200), nullable=True)
    industry = db.Column(db.String(100), nullable=True)
    website = db.Column(db.String(200), nullable=True)
    
    # Descriptions
    company_description = db.Column(db.Text, nullable=True)
    about_company = db.Column(db.Text, nullable=True)
    products_services = db.Column(db.Text, nullable=True)
    
    # Team & Culture
    team_info = db.Column(db.Text, nullable=True)
    company_culture = db.Column(db.Text, nullable=True)
    
    # Contact
    contact_info = db.Column(db.Text, nullable=True)
    custom_knowledge = db.Column(db.Text, nullable=True)
    
    # Metadata
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    def __repr__(self):
        return f'<CompanyInfo {self.company_name}>'
    
    def get_full_knowledge(self):
        """Return all company knowledge as formatted text"""
        text = ""
        
        if self.company_name:
            text += f"# Company: {self.company_name}\n\n"
        
        if self.company_description:
            text += f"## Description\n{self.company_description}\n\n"
        
        if self.about_company:
            text += f"## About Us\n{self.about_company}\n\n"
        
        if self.products_services:
            text += f"## Products & Services\n{self.products_services}\n\n"
        
        if self.team_info:
            text += f"## Team\n{self.team_info}\n\n"
        
        if self.company_culture:
            text += f"## Culture\n{self.company_culture}\n\n"
        
        if self.contact_info:
            text += f"## Contact\n{self.contact_info}\n\n"
        
        if self.custom_knowledge:
            text += f"## Additional Knowledge\n{self.custom_knowledge}\n\n"
        
        return text
    
    @staticmethod
    def get_or_create():
        """Get or create the company info entry"""
        company = CompanyInfo.query.first()
        if not company:
            company = CompanyInfo()
            db.session.add(company)
            db.session.commit()
        return company


class Knowledge(db.Model):
    """Model for storing custom knowledge entries"""
    __tablename__ = 'knowledge'
    
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), nullable=False)
    category = db.Column(db.String(100), nullable=True)
    content = db.Column(db.Text, nullable=False)
    description = db.Column(db.String(500), nullable=True)
    
    # Metadata
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    created_by = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    
    # Status
    is_active = db.Column(db.Boolean, default=True)
    
    # Relationships
    creator = db.relationship('User', backref='knowledge_created')
    
    def __repr__(self):
        return f'<Knowledge {self.title}>'
    
    def to_dict(self):
        """Convert to dictionary for JSON responses"""
        return {
            'id': self.id,
            'title': self.title,
            'category': self.category,
            'content': self.content,
            'description': self.description,
            'created_at': self.created_at.strftime('%Y-%m-%d %H:%M:%S'),
            'updated_at': self.updated_at.strftime('%Y-%m-%d %H:%M:%S'),
            'is_active': self.is_active
        }
    
    @staticmethod
    def get_all_active_knowledge():
        """Get all active knowledge entries formatted for AI context"""
        knowledge_entries = Knowledge.query.filter_by(is_active=True).all()
        
        if not knowledge_entries:
            return ""
        
        knowledge_text = "# CUSTOM KNOWLEDGE BASE\n\n"
        
        # Group by category
        by_category = {}
        for entry in knowledge_entries:
            cat = entry.category or "General"
            if cat not in by_category:
                by_category[cat] = []
            by_category[cat].append(entry)
        
        # Format knowledge entries
        for category, entries in by_category.items():
            knowledge_text += f"## {category}\n\n"
            for entry in entries:
                knowledge_text += f"### {entry.title}\n"
                if entry.description:
                    knowledge_text += f"{entry.description}\n\n"
                knowledge_text += f"{entry.content}\n\n"
        
        return knowledge_text
    
    @staticmethod
    def get_knowledge_by_category(category):
        """Get knowledge by specific category"""
        entries = Knowledge.query.filter_by(category=category, is_active=True).all()
        return entries
    
    @staticmethod
    def search_knowledge(query_text):
        """Search knowledge entries"""
        return Knowledge.query.filter(
            (Knowledge.title.ilike(f'%{query_text}%')) |
            (Knowledge.content.ilike(f'%{query_text}%')) |
            (Knowledge.description.ilike(f'%{query_text}%')),
            Knowledge.is_active == True
        ).all()


# ============================================================
# RESPONSE FORMATTING UTILITIES
# ============================================================

def format_response_with_table(query_text, ai_response):
    """
    Check if response should include a table, and format accordingly
    
    Args:
        query_text: The original query
        ai_response: The AI-generated response
    
    Returns:
        Formatted response with embedded HTML table if applicable
    """
    
    detector = TableDetector(confidence_threshold=0.6)
    is_comparison, confidence = detector.detect_comparison_question(query_text)
    
    if not is_comparison or confidence < 0.6:
        return ai_response
    
    print(f"[FORMAT] Creating table for comparison query")
    
    # Extract comparison items from response (simplified version)
    # In production, use more sophisticated parsing
    # Check if response contains typical table-like structure
    has_table_structure = any(pattern in ai_response.lower() for pattern in [
        'vs', 'versus', 'compared to', 'difference', 'same', 'unlike'
    ])
    
    if has_table_structure:
        # Add visual separator and note about table
        formatted = f"""
{ai_response}

---

💡 **Note:** This comparison answer contains information that could be displayed in a table format. The AI has provided a detailed comparison above. For a quick side-by-side view, you can export this as a PDF or ask for a specific table format.
        """.strip()
        return formatted
    
    return ai_response


def extract_and_generate_table(query_text, ai_response, table_type='comparison'):
    """
    Extract structured data from AI response and generate HTML table
    
    Args:
        query_text: Original question
        ai_response: AI's response text
        table_type: 'comparison' or 'pros_cons'
    
    Returns:
        Tuple of (formatted_response, html_table)
    """
    
    # This is a simplified example
    # In production, you'd parse the AI response more intelligently
    
    generator = HTMLTableGenerator()
    
    if table_type == 'comparison':
        # Example: Extract items and attributes
        # You'd parse ai_response to get this data
        
        example_data = {
            "Item A": {"Feature 1": "Value", "Feature 2": "Value"},
            "Item B": {"Feature 1": "Value", "Feature 2": "Value"},
        }
        
        table_html = generator.create_comparison_table(
            title="Comparison",
            items=["Item A", "Item B"],
            attributes=["Feature 1", "Feature 2"],
            data=example_data,
            theme="beige"
        )
        
        return ai_response, table_html
    
    elif table_type == 'pros_cons':
        # Example for pros/cons
        table_html = generator.create_pros_cons_table(
            item1="Option A",
            pros1=["Pro 1", "Pro 2"],
            cons1=["Con 1", "Con 2"],
            item2="Option B",
            pros2=["Pro 1", "Pro 2"],
            cons2=["Con 1", "Con 2"]
        )
        
        return ai_response, table_html
    
    return ai_response, None


# ============================================================
# AUTH MIDDLEWARE & DECORATORS
# ============================================================

@app.before_request
def make_session_permanent():
    """Make session permanent"""
    session.permanent = True
    app.permanent_session_lifetime = timedelta(days=7)


def login_required(f):
    """Decorator to require login"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            flash('Please log in first', 'error')
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function


def admin_required(f):
    """Decorator to require admin access"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            flash('Please log in first', 'error')
            return redirect(url_for('login'))
        
        user = User.query.get(session['user_id'])
        if not user or not user.is_admin:
            flash('Admin access required', 'error')
            return redirect(url_for('index'))
        return f(*args, **kwargs)
    return decorated_function


# ============================================================
# AUTHENTICATION ROUTES
# ============================================================

@app.route('/register', methods=['GET', 'POST'])
def register():
    """User registration"""
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '')
        confirm_password = request.form.get('confirm_password', '')
        
        if not username or not password:
            flash('Username and password are required', 'error')
            return render_template('register.html')
        
        if password != confirm_password:
            flash('Passwords do not match', 'error')
            return render_template('register.html')
        
        if len(password) < 6:
            flash('Password must be at least 6 characters', 'error')
            return render_template('register.html')
        
        if User.query.filter_by(username=username).first():
            flash('Username already exists', 'error')
            return render_template('register.html')
        
        user = User(username=username)
        user.set_password(password)
        db.session.add(user)
        db.session.commit()
        
        flash('Account created successfully! Please login.', 'success')
        return redirect(url_for('login'))
    
    return render_template('register.html')


@app.route('/login', methods=['GET', 'POST'])
def login():
    """User login"""
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '')
        
        user = User.query.filter_by(username=username).first()
        
        if user and user.check_password(password):
            session['user_id'] = user.id
            session['username'] = user.username
            session['is_admin'] = user.is_admin
            flash(f'Welcome back, {user.username}!', 'success')
            
            if user.is_admin:
                return redirect(url_for('admin_dashboard'))
            return redirect(url_for('index'))
        else:
            flash('Invalid username or password', 'error')
    
    return render_template('login.html')


@app.route('/logout')
@login_required
def logout():
    """User logout"""
    session.clear()
    flash('You have been logged out', 'success')
    return redirect(url_for('login'))


# ============================================================
# MAIN ROUTES
# ============================================================

@app.route('/')
@login_required
def index():
    """Home page"""
    user = User.query.get(session['user_id'])
    queries = Query.query.filter_by(user_id=user.id).order_by(Query.created_at.desc()).limit(5).all()
    return render_template('index.html', queries=queries, user=user)


@app.route('/history')
@login_required
def history():
    """Query history page"""
    user = User.query.get(session['user_id'])
    page = request.args.get('page', 1, type=int)
    
    queries = Query.query.filter_by(user_id=user.id)\
        .order_by(Query.created_at.desc())\
        .paginate(page=page, per_page=10)
    
    return render_template('history.html', queries=queries, user=user)


@app.route('/query/<int:query_id>')
@login_required
def view_query(query_id):
    """View query details"""
    query = Query.query.get_or_404(query_id)
    
    # Check ownership
    if query.user_id != session['user_id']:
        flash('You do not have access to this query', 'error')
        return redirect(url_for('history'))
    
    return render_template('query_detail.html', query=query)


@app.route('/settings')
@login_required
def settings():
    """User settings page"""
    user = User.query.get(session['user_id'])
    return render_template('settings.html', user=user)


# ============================================================
# ADMIN ROUTES
# ============================================================

@app.route('/admin')
@admin_required
def admin_dashboard():
    """Admin dashboard"""
    total_users = User.query.count()
    total_queries = Query.query.count()
    completed_queries = Query.query.filter_by(status='completed').count()
    failed_queries = Query.query.filter_by(status='failed').count()
    
    # Task type breakdown
    task_breakdown = {}
    for query in Query.query.all():
        task = query.task_type
        task_breakdown[task] = task_breakdown.get(task, 0) + 1
    
    # Recent queries
    recent_queries = Query.query.order_by(Query.created_at.desc()).limit(5).all()
    
    # Knowledge statistics
    knowledge_total = Knowledge.query.count()
    knowledge_active = Knowledge.query.filter_by(is_active=True).count()
    
    # Company knowledge status
    company_info = CompanyInfo.query.first()
    company_knowledge_status = 'Configured' if company_info and company_info.company_name else 'Not Configured'
    company_name = company_info.company_name if company_info else 'Not Set'
    
    return render_template('admin/dashboard.html', 
        total_users=total_users,
        total_queries=total_queries,
        completed_queries=completed_queries,
        failed_queries=failed_queries,
        task_breakdown=task_breakdown,
        recent_queries=recent_queries,
        knowledge_total=knowledge_total,
        knowledge_active=knowledge_active,
        company_knowledge_status=company_knowledge_status,
        company_name=company_name)


@app.route('/admin/users')
@admin_required
def admin_users():
    """Manage users"""
    page = request.args.get('page', 1, type=int)
    users = User.query.paginate(page=page, per_page=10)
    return render_template('admin/users.html', users=users)


@app.route('/admin/queries')
@admin_required
def admin_queries():
    """Manage queries"""
    page = request.args.get('page', 1, type=int)
    queries = Query.query.order_by(Query.created_at.desc()).paginate(page=page, per_page=10)
    return render_template('admin/queries.html', queries=queries)


@app.route('/admin/user/<int:user_id>/queries')
@admin_required
def admin_user_queries(user_id):
    """View user's queries"""
    user = User.query.get_or_404(user_id)
    page = request.args.get('page', 1, type=int)
    queries = Query.query.filter_by(user_id=user_id).order_by(Query.created_at.desc()).paginate(page=page, per_page=10)
    return render_template('admin/user_queries.html', user=user, queries=queries)


# ============================================================
# COMPANY INFO ROUTES
# ============================================================

@app.route('/admin/company-info', methods=['GET', 'POST'])
@admin_required
def manage_company_info():
    """Manage company information"""
    company_info = CompanyInfo.query.first()
    
    if request.method == 'POST':
        data = request.form
        
        if not company_info:
            company_info = CompanyInfo()
        
        company_info.company_name = data.get('company_name', '')
        company_info.company_description = data.get('company_description', '')
        company_info.about_company = data.get('about_company', '')
        company_info.products_services = data.get('products_services', '')
        company_info.company_culture = data.get('company_culture', '')
        company_info.team_info = data.get('team_info', '')
        company_info.contact_info = data.get('contact_info', '')
        company_info.website = data.get('website', '')
        company_info.industry = data.get('industry', '')
        company_info.custom_knowledge = data.get('custom_knowledge', '')
        company_info.updated_at = datetime.utcnow()
        
        db.session.add(company_info)
        db.session.commit()
        
        flash('Company information updated successfully!', 'success')
        return redirect(url_for('manage_company_info'))
    
    return render_template('admin/company_info.html', company_info=company_info)


@app.route('/api/get-company-knowledge')
def get_company_knowledge():
    """Get company knowledge for use in queries"""
    company_info = CompanyInfo.query.first()
    
    if company_info:
        return jsonify({
            'knowledge': company_info.get_full_knowledge(),
            'company_name': company_info.company_name
        })
    else:
        return jsonify({
            'knowledge': '',
            'company_name': ''
        })


@app.route('/api/company-knowledge/export')
@admin_required
def export_company_knowledge():
    """Export company knowledge as JSON"""
    company_info = CompanyInfo.query.first()
    
    if not company_info or not company_info.company_name:
        return jsonify({'error': 'No company knowledge configured'}), 404
    
    export_data = {
        'company_name': company_info.company_name,
        'company_description': company_info.company_description,
        'about_company': company_info.about_company,
        'products_services': company_info.products_services,
        'company_culture': company_info.company_culture,
        'team_info': company_info.team_info,
        'contact_info': company_info.contact_info,
        'website': company_info.website,
        'industry': company_info.industry,
        'custom_knowledge': company_info.custom_knowledge,
        'exported_at': company_info.updated_at.isoformat()
    }
    
    return send_file(
        io.BytesIO(json.dumps(export_data, indent=2).encode('utf-8')),
        mimetype='application/json',
        as_attachment=True,
        download_name=f"{company_info.company_name.replace(' ', '_')}_knowledge.json"
    )


@app.route('/api/company-knowledge/status')
def company_knowledge_status():
    """Get status of company knowledge setup"""
    company_info = CompanyInfo.query.first()
    
    if not company_info:
        return jsonify({
            'configured': False,
            'company_name': None,
            'completeness': 0
        })
    
    # Calculate completeness percentage
    fields = [
        company_info.company_name,
        company_info.company_description,
        company_info.about_company,
        company_info.products_services,
        company_info.company_culture,
        company_info.team_info,
        company_info.contact_info,
        company_info.website,
        company_info.industry,
        company_info.custom_knowledge
    ]
    
    filled_fields = sum(1 for field in fields if field)
    completeness = int((filled_fields / len(fields)) * 100)
    
    return jsonify({
        'configured': bool(company_info.company_name),
        'company_name': company_info.company_name,
        'completeness': completeness,
        'fields_filled': filled_fields,
        'total_fields': len(fields),
        'last_updated': company_info.updated_at.isoformat()
    })


@app.route('/api/company-knowledge/search', methods=['POST'])
@login_required
def search_company_knowledge():
    """Search within company knowledge"""
    data = request.get_json()
    search_term = data.get('search', '').lower().strip()
    
    if not search_term or len(search_term) < 2:
        return jsonify({'error': 'Search term must be at least 2 characters'}), 400
    
    company_info = CompanyInfo.query.first()
    
    if not company_info:
        return jsonify({'results': []})
    
    results = {}
    
    # Search in each field
    if company_info.company_name and search_term in company_info.company_name.lower():
        results['company_name'] = company_info.company_name
    
    if company_info.company_description and search_term in company_info.company_description.lower():
        results['company_description'] = company_info.company_description
    
    if company_info.about_company and search_term in company_info.about_company.lower():
        results['about_company'] = company_info.about_company
    
    if company_info.products_services and search_term in company_info.products_services.lower():
        results['products_services'] = company_info.products_services
    
    if company_info.company_culture and search_term in company_info.company_culture.lower():
        results['company_culture'] = company_info.company_culture
    
    if company_info.team_info and search_term in company_info.team_info.lower():
        results['team_info'] = company_info.team_info
    
    if company_info.contact_info and search_term in company_info.contact_info.lower():
        results['contact_info'] = company_info.contact_info
    
    if company_info.custom_knowledge and search_term in company_info.custom_knowledge.lower():
        results['custom_knowledge'] = company_info.custom_knowledge
    
    return jsonify({
        'search_term': search_term,
        'results': results,
        'found_in_fields': len(results)
    })


# ============================================================
# QUERY EXECUTION ROUTES
# ============================================================

def detect_task_type(query_text):
    """Detect task type from query text"""
    query_lower = query_text.lower()
    
    if any(word in query_lower for word in ['code', 'python', 'javascript', 'function', 'algorithm']):
        return 'code'
    elif any(word in query_lower for word in ['analyze', 'analysis', 'explain', 'compare']):
        return 'analysis'
    elif any(word in query_lower for word in ['write', 'story', 'poem', 'create', 'generate']):
        return 'creative'
    elif any(word in query_lower for word in ['research', 'study', 'find', 'investigate']):
        return 'research'
    elif any(word in query_lower for word in ['solve', 'problem', 'fix', 'debug']):
        return 'problem_solving'
    else:
        return 'general'


def get_knowledge_context():
    """Get all active knowledge as context for AI queries"""
    # Get company knowledge
    company = CompanyInfo.get_or_create()
    company_knowledge = company.get_full_knowledge()
    
    # Get custom knowledge
    custom_knowledge = Knowledge.get_all_active_knowledge()
    
    # Combine both
    full_context = company_knowledge + "\n\n" + custom_knowledge
    
    return full_context.strip() if full_context.strip() else ""


@app.route('/api/detect-table-type', methods=['POST'])
@login_required
def detect_table_type():
    """
    API endpoint to check if a query should generate a table
    Call this BEFORE executing the query
    """
    data = request.get_json()
    query_text = data.get('query', '')
    
    detector = TableDetector()
    is_comparison, confidence = detector.detect_comparison_question(query_text)
    keywords = detector.get_detected_keywords(query_text)
    
    return jsonify({
        'should_generate_table': is_comparison and confidence >= 0.6,
        'confidence': confidence,
        'detected_keywords': keywords,
        'query': query_text
    })


@app.route('/api/execute-query', methods=['POST'])
@login_required
def execute_query():
    """Updated execute_query with table detection"""
    
    data = request.get_json()
    query_text = data.get('query', '').strip()
    
    if not query_text:
        return jsonify({'error': 'Query cannot be empty'}), 400
    
    # Detect if this is a comparison question
    detector = TableDetector()
    is_comparison, confidence = detector.detect_comparison_question(query_text)
    
    # Create query record
    query = Query(
        user_id=session['user_id'],
        query_text=query_text,
        task_type=detect_task_type(query_text),
        status='processing',
        is_comparison_query=is_comparison,
        comparison_confidence=confidence
    )
    
    db.session.add(query)
    db.session.commit()
    
    # Execute in background (your existing code)
    from threading import Thread
    thread = Thread(
        target=execute_query_background,
        args=(query.id, query_text, session['user_id'])
    )
    thread.daemon = True
    thread.start()
    
    return jsonify({
        'query_id': query.id,
        'detected_as_comparison': is_comparison,
        'confidence': confidence
    })


def execute_query_background(query_id, query_text=None, user_id=None):
    """Execute query in background using LLM"""
    from openai import OpenAI
    
    try:
        with app.app_context():
            query = Query.query.get(query_id)
            start_time = time.time()
            
            # ============================================================
            # TABLE DETECTION PHASE
            # ============================================================
            # Initialize detector for comparison/difference questions
            detector = TableDetector(confidence_threshold=0.6)
            
            # Check if this is a comparison/difference question
            is_comparison, confidence = detector.detect_comparison_question(query.query_text)
            
            print(f"[TABLE DETECTION] Query: {query.query_text[:80]}")
            print(f"[TABLE DETECTION] Is Comparison: {is_comparison}, Confidence: {confidence:.2f}")
            
            # Store detection results in database
            query.is_comparison_query = is_comparison
            query.comparison_confidence = confidence
            
            if is_comparison and confidence >= 0.6:
                print(f"[TABLE GENERATION] This is a comparison question - will generate table")
                detected_keywords = detector.get_detected_keywords(query.query_text)
                print(f"[TABLE DETECTION] Keywords detected: {detected_keywords}")
            
            # ============================================================
            # LLM PROCESSING PHASE
            # ============================================================
            
            # Initialize LLM
            openrouter_client = OpenAI(
                base_url="https://openrouter.ai/api/v1",
                api_key=os.getenv('OPENROUTER_API_KEY', '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'),
            )
            
            # Get knowledge context
            knowledge_context = get_knowledge_context()
            
            # System prompt with knowledge integration
            system_instruction = """You are an expert research assistant with advanced reasoning capabilities.

Your task is to:
1. Thoroughly analyze the question
2. Identify key aspects and subtopics
3. Provide comprehensive, well-reasoned responses
4. Consider multiple perspectives and evidence
5. Support claims with specific examples

Format your response clearly with sections and examples."""
            
            # Add knowledge to system prompt if available
            if knowledge_context:
                system_instruction += f"\n\n--- KNOWLEDGE BASE ---\n{knowledge_context}\n--- END KNOWLEDGE BASE ---\n\nUse this knowledge base to answer questions accurately when relevant."
            
            # Build messages
            messages = [
                {
                    "role": "system",
                    "content": system_instruction
                },
                {
                    "role": "user",
                    "content": query.query_text
                }
            ]
            
            # Execute query with reasoning
            response = openrouter_client.chat.completions.create(
                model="openai/gpt-oss-20b:free",
                messages=messages,
                max_tokens=8000,
                extra_body={"reasoning": {"enabled": True}}
            )
            
            assistant_message = response.choices[0].message
            response_content = assistant_message.content
            reasoning_raw = getattr(assistant_message, 'reasoning_details', '')
            
            # Convert reasoning to string (handle list or dict formats)
            if isinstance(reasoning_raw, list):
                reasoning_text = '\n'.join([
                    item.get('text', str(item)) if isinstance(item, dict) else str(item)
                    for item in reasoning_raw
                ])
            elif isinstance(reasoning_raw, dict):
                reasoning_text = reasoning_raw.get('text', json.dumps(reasoning_raw))
            else:
                reasoning_text = str(reasoning_raw) if reasoning_raw else ''
            
            execution_time = time.time() - start_time
            
            # ============================================================
            # TABLE GENERATION PHASE
            # ============================================================
            # Generate table if this is a comparison question
            table_html = None
            if query.is_comparison_query and query.comparison_confidence >= 0.6:
                try:
                    print(f"[TABLE GENERATION] Generating HTML table for comparison query")
                    
                    # For now, we store the detection info
                    # In production, you could parse the response to extract structured data
                    # and generate a detailed comparison table
                    # 
                    # Example of advanced table generation:
                    # table_html = HTMLTableGenerator.create_comparison_table(
                    #     title=f"Comparison Analysis",
                    #     items=[...],  # Extract from response
                    #     attributes=[...],  # Extract from response
                    #     data={...},  # Extract from response
                    #     theme="beige"
                    # )
                    
                    print(f"[TABLE GENERATION] Ready to generate table (confidence: {query.comparison_confidence:.2f})")
                except Exception as e:
                    print(f"[TABLE GENERATION ERROR] {str(e)}")
            
            # ============================================================
            # UPDATE QUERY IN DATABASE
            # ============================================================
            # Update query
            query.status = 'completed'
            query.response = response_content
            query.reasoning = reasoning_text
            query.execution_time = execution_time
            query.tools_used = 'search, research'
            
            # Store table HTML if generated
            if table_html:
                query.table_html = table_html
                print(f"[TABLE STORAGE] Table HTML stored in database")
            
            db.session.commit()
            
            print(f"[QUERY COMPLETE] Query {query_id} completed in {execution_time:.2f}s")
            print(f"[COMPARISON DETECTION] Stored: is_comparison={query.is_comparison_query}, confidence={query.comparison_confidence:.2f}")
    
    except Exception as e:
        with app.app_context():
            query = Query.query.get(query_id)
            query.status = 'failed'
            query.error_message = str(e)
            
            # Ensure comparison fields are set even on error
            if not query.is_comparison_query:
                detector = TableDetector()
                is_comparison, confidence = detector.detect_comparison_question(query.query_text)
                query.is_comparison_query = is_comparison
                query.comparison_confidence = confidence
            
            print(f"[QUERY ERROR] Query {query_id} failed: {str(e)}")
            db.session.commit()


@app.route('/api/query-status/<int:query_id>')
@login_required
def query_status(query_id):
    """Get query status"""
    query = Query.query.get_or_404(query_id)
    
    if query.user_id != session['user_id']:
        return jsonify({'error': 'Unauthorized'}), 403
    
    return jsonify({
        'status': query.status,
        'query_id': query.id
    })


@app.route('/api/statistics')
@login_required
def get_statistics():
    """Get user statistics"""
    user = User.query.get(session['user_id'])
    
    total_queries = Query.query.filter_by(user_id=user.id).count()
    completed_queries = Query.query.filter_by(user_id=user.id, status='completed').count()
    failed_queries = Query.query.filter_by(user_id=user.id, status='failed').count()
    
    task_breakdown = {}
    for query in Query.query.filter_by(user_id=user.id).all():
        task = query.task_type
        task_breakdown[task] = task_breakdown.get(task, 0) + 1
    
    return jsonify({
        'total_queries': total_queries,
        'completed_queries': completed_queries,
        'failed_queries': failed_queries,
        'task_breakdown': task_breakdown
    })


@app.route('/api/change-password', methods=['POST'])
@login_required
def change_password():
    """Change user password"""
    user = User.query.get(session['user_id'])
    data = request.get_json()
    
    current_password = data.get('currentPassword', '')
    new_password = data.get('newPassword', '')
    
    if not user.check_password(current_password):
        return jsonify({'error': 'Current password is incorrect'}), 400
    
    if len(new_password) < 6:
        return jsonify({'error': 'New password must be at least 6 characters'}), 400
    
    user.set_password(new_password)
    db.session.commit()
    
    return jsonify({'success': True})


# ============================================================
# CUSTOM KNOWLEDGE ROUTES
# ============================================================

@app.route('/admin/knowledge', methods=['GET'])
@login_required
@admin_required
def manage_knowledge():
    """Display all knowledge entries with pagination and search"""
    page = request.args.get('page', 1, type=int)
    search_query = request.args.get('search', '', type=str)
    category_filter = request.args.get('category', '', type=str)
    
    query = Knowledge.query
    
    # Apply search filter
    if search_query:
        query = query.filter(
            (Knowledge.title.ilike(f'%{search_query}%')) |
            (Knowledge.content.ilike(f'%{search_query}%'))
        )
    
    # Apply category filter
    if category_filter:
        query = query.filter_by(category=category_filter)
    
    # Paginate results
    knowledge_paginated = query.order_by(Knowledge.created_at.desc()).paginate(
        page=page, per_page=10
    )
    
    # Get all categories for filter dropdown
    categories = db.session.query(Knowledge.category).distinct().filter(
        Knowledge.category.isnot(None)
    ).all()
    categories = [cat[0] for cat in categories]
    
    return render_template('admin/knowledge.html',
                         knowledge=knowledge_paginated,
                         search_query=search_query,
                         categories=categories,
                         selected_category=category_filter)


@app.route('/admin/knowledge/create', methods=['GET', 'POST'])
@login_required
@admin_required
def create_knowledge():
    """Create new knowledge entry"""
    if request.method == 'POST':
        title = request.form.get('title', '').strip()
        category = request.form.get('category', '').strip() or None
        content = request.form.get('content', '').strip()
        description = request.form.get('description', '').strip() or None
        
        # Validation
        if not title or not content:
            flash('Title and content are required', 'error')
            return redirect(url_for('create_knowledge'))
        
        try:
            # Create new knowledge entry
            knowledge = Knowledge(
                title=title,
                category=category,
                content=content,
                description=description,
                created_by=session['user_id'],
                is_active=True
            )
            
            db.session.add(knowledge)
            db.session.commit()
            
            flash(f'Knowledge "{title}" created successfully!', 'success')
            return redirect(url_for('manage_knowledge'))
        
        except Exception as e:
            db.session.rollback()
            flash(f'Error creating knowledge: {str(e)}', 'error')
            return redirect(url_for('create_knowledge'))
    
    return render_template('admin/knowledge_form.html', knowledge=None)


@app.route('/admin/knowledge/<int:knowledge_id>/edit', methods=['GET', 'POST'])
@login_required
@admin_required
def edit_knowledge(knowledge_id):
    """Edit existing knowledge entry"""
    knowledge = Knowledge.query.get_or_404(knowledge_id)
    
    if request.method == 'POST':
        title = request.form.get('title', '').strip()
        category = request.form.get('category', '').strip() or None
        content = request.form.get('content', '').strip()
        description = request.form.get('description', '').strip() or None
        is_active = request.form.get('is_active') == 'on'
        
        # Validation
        if not title or not content:
            flash('Title and content are required', 'error')
            return redirect(url_for('edit_knowledge', knowledge_id=knowledge_id))
        
        try:
            knowledge.title = title
            knowledge.category = category
            knowledge.content = content
            knowledge.description = description
            knowledge.is_active = is_active
            knowledge.updated_at = datetime.utcnow()
            
            db.session.commit()
            
            flash(f'Knowledge "{title}" updated successfully!', 'success')
            return redirect(url_for('manage_knowledge'))
        
        except Exception as e:
            db.session.rollback()
            flash(f'Error updating knowledge: {str(e)}', 'error')
            return redirect(url_for('edit_knowledge', knowledge_id=knowledge_id))
    
    return render_template('admin/knowledge_form.html', knowledge=knowledge)


@app.route('/admin/knowledge/<int:knowledge_id>/delete', methods=['POST'])
@login_required
@admin_required
def delete_knowledge(knowledge_id):
    """Delete knowledge entry"""
    knowledge = Knowledge.query.get_or_404(knowledge_id)
    title = knowledge.title
    
    try:
        db.session.delete(knowledge)
        db.session.commit()
        flash(f'Knowledge "{title}" deleted successfully!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error deleting knowledge: {str(e)}', 'error')
    
    return redirect(url_for('manage_knowledge'))


@app.route('/admin/knowledge/<int:knowledge_id>/toggle', methods=['POST'])
@login_required
@admin_required
def toggle_knowledge_active(knowledge_id):
    """Toggle knowledge entry active status"""
    knowledge = Knowledge.query.get_or_404(knowledge_id)
    
    try:
        knowledge.is_active = not knowledge.is_active
        db.session.commit()
        
        status = "activated" if knowledge.is_active else "deactivated"
        flash(f'Knowledge "{knowledge.title}" {status}!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error toggling knowledge: {str(e)}', 'error')
    
    return redirect(url_for('manage_knowledge'))


@app.route('/api/knowledge/search', methods=['GET'])
@login_required
def search_knowledge_api():
    """API endpoint to search knowledge"""
    search_query = request.args.get('q', '', type=str)
    
    if len(search_query) < 2:
        return jsonify([])
    
    results = Knowledge.search_knowledge(search_query)
    
    return jsonify([{
        'id': r.id,
        'title': r.title,
        'category': r.category,
        'snippet': r.content[:100] + '...' if len(r.content) > 100 else r.content
    } for r in results[:10]])


@app.route('/api/knowledge/categories', methods=['GET'])
@login_required
@admin_required
def get_knowledge_categories():
    """Get all knowledge categories"""
    categories = db.session.query(Knowledge.category).distinct().filter(
        Knowledge.category.isnot(None)
    ).all()
    categories = [cat[0] for cat in categories]
    return jsonify(categories)


@app.route('/api/knowledge/stats', methods=['GET'])
@login_required
@admin_required
def get_knowledge_stats():
    """Get knowledge statistics for dashboard"""
    total = Knowledge.query.count()
    active = Knowledge.query.filter_by(is_active=True).count()
    inactive = Knowledge.query.filter_by(is_active=False).count()
    
    # Count by category
    by_category = db.session.query(
        Knowledge.category,
        db.func.count(Knowledge.id)
    ).group_by(Knowledge.category).all()
    
    category_stats = {cat or 'Uncategorized': count for cat, count in by_category}
    
    return jsonify({
        'total': total,
        'active': active,
        'inactive': inactive,
        'by_category': category_stats
    })


# ============================================================
# EXPORT ROUTES
# ============================================================

@app.route('/export/query/<int:query_id>/<format>')
@login_required
def export_query(query_id, format):
    """Export query to different formats"""
    query = Query.query.get_or_404(query_id)
    
    if query.user_id != session['user_id']:
        flash('You do not have access to this query', 'error')
        return redirect(url_for('history'))
    
    if format == 'txt':
        return export_to_txt(query)
    elif format == 'pdf':
        return export_to_pdf(query)
    elif format == 'doc':
        return export_to_doc(query)
    else:
        flash('Invalid export format', 'error')
        return redirect(url_for('view_query', query_id=query_id))


def export_to_txt(query):
    """Export query to TXT format"""
    execution_str = f"{query.execution_time:.2f} seconds" if query.execution_time else "N/A"
    
    content = f"""AI RESEARCH AGENT - QUERY EXPORT
=====================================

Query: {query.query_text}
Task Type: {query.task_type}
Status: {query.status}
Created: {query.created_at.strftime('%Y-%m-%d %H:%M:%S')}
Execution Time: {execution_str}

=====================================
RESPONSE
=====================================
{query.response or 'No response available'}

=====================================
REASONING PROCESS
=====================================
{query.reasoning or 'No reasoning available'}

=====================================
TOOLS USED
=====================================
{query.tools_used or 'None'}
"""
    
    return send_file(
        io.BytesIO(content.encode('utf-8')),
        mimetype='text/plain',
        as_attachment=True,
        download_name=f'query_{query.id}.txt'
    )


def export_to_pdf(query):
    """Export query to PDF format"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    story = []
    styles = getSampleStyleSheet()
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
    )
    story.append(Paragraph("AI Research Agent - Query Export", title_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Query info
    info_style = styles['Normal']
    execution_str = f"{query.execution_time:.2f} seconds" if query.execution_time else "N/A"
    
    story.append(Paragraph(f"<b>Query:</b> {query.query_text}", info_style))
    story.append(Paragraph(f"<b>Task Type:</b> {query.task_type}", info_style))
    story.append(Paragraph(f"<b>Status:</b> {query.status}", info_style))
    story.append(Paragraph(f"<b>Created:</b> {query.created_at.strftime('%Y-%m-%d %H:%M:%S')}", info_style))
    story.append(Paragraph(f"<b>Execution Time:</b> {execution_str}", info_style))
    story.append(Spacer(1, 0.3*inch))
    
    # Response
    story.append(Paragraph("<b>Response</b>", styles['Heading2']))
    story.append(Paragraph(query.response or "No response available", info_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Reasoning
    if query.reasoning:
        story.append(PageBreak())
        story.append(Paragraph("<b>Reasoning Process</b>", styles['Heading2']))
        story.append(Paragraph(query.reasoning, info_style))
    
    # Tools
    if query.tools_used:
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(f"<b>Tools Used:</b> {query.tools_used}", info_style))
    
    doc.build(story)
    buffer.seek(0)
    
    return send_file(
        buffer,
        mimetype='application/pdf',
        as_attachment=True,
        download_name=f'query_{query.id}.pdf'
    )


def export_to_doc(query):
    """Export query to DOC format (using docx)"""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_paragraph("AI Research Agent - Query Export")
        title_format = title.paragraph_format
        title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(24)
        title.runs[0].font.bold = True
        
        # Query info
        doc.add_paragraph(f"Query: {query.query_text}")
        doc.add_paragraph(f"Task Type: {query.task_type}")
        doc.add_paragraph(f"Status: {query.status}")
        doc.add_paragraph(f"Created: {query.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
        if query.execution_time:
            doc.add_paragraph(f"Execution Time: {query.execution_time:.2f} seconds")
        
        doc.add_paragraph()
        
        # Response
        response_heading = doc.add_paragraph("Response")
        response_heading.runs[0].font.bold = True
        response_heading.runs[0].font.size = Pt(14)
        doc.add_paragraph(query.response or "No response available")
        
        # Reasoning
        if query.reasoning:
            doc.add_page_break()
            reasoning_heading = doc.add_paragraph("Reasoning Process")
            reasoning_heading.runs[0].font.bold = True
            reasoning_heading.runs[0].font.size = Pt(14)
            doc.add_paragraph(query.reasoning)
        
        # Tools
        if query.tools_used:
            doc.add_paragraph(f"Tools Used: {query.tools_used}")
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'query_{query.id}.docx'
        )
    except ImportError:
        flash('python-docx is required for DOC export. Please install it.', 'error')
        return redirect(url_for('view_query', query_id=query.id))


# ============================================================
# ERROR HANDLERS
# ============================================================

@app.errorhandler(404)
def not_found(error):
    """404 error handler"""
    return render_template('error.html', error='Page not found'), 404


@app.errorhandler(500)
def server_error(error):
    """500 error handler"""
    return render_template('error.html', error='Server error'), 500


# ============================================================
# DATABASE INITIALIZATION
# ============================================================

def init_db():
    """Initialize database"""
    with app.app_context():
        db.create_all()
        
        # Create admin user if doesn't exist
        admin = User.query.filter_by(username='admin').first()
        if not admin:
            admin = User(username='admin', is_admin=True)
            admin.set_password('admin123')
            db.session.add(admin)
            db.session.commit()
            print("✅ Admin user created: username=admin, password=admin123")


if __name__ == '__main__':
    init_db()
    app.run(debug=True, port=5000)